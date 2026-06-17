from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = r"D:\Desktop\develop\Android Studio\DayMarkApp\docs\每日打卡习惯追踪APP项目报告.docx"


def set_run_font(run, size=10.5, bold=False):
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)


def set_style_font(style):
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

normal = doc.styles["Normal"]
set_style_font(normal)
normal.font.size = Pt(10.5)

for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
    set_style_font(doc.styles[style_name])


def add_paragraph(text="", *, size=10.5, bold=False, align=None, first_line_cm=0.74):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.first_line_indent = Cm(first_line_cm) if first_line_cm else Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 11}.get(level, 11), bold=True)
    return p


def fill_cell(cell, text, *, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
cover.paragraph_format.space_after = Pt(18)
run = cover.add_run("每日打卡习惯追踪APP\n项目报告")
set_run_font(run, size=22, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("DayMark 项目扫描与课程设计成稿")
set_run_font(run, size=13)

add_paragraph("项目名称：DayMark 每日打卡习惯追踪 APP", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=0, size=12)
add_paragraph("开发平台：Android Studio Panda 2 | 2025.3.2", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=0, size=12)
add_paragraph("开发语言：Java + XML", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=0, size=12)
add_paragraph("姓名：__________    学号：__________    班级：__________", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=0, size=12)
add_paragraph("生成日期：2026年6月17日", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=0, size=12)

doc.add_page_break()

add_paragraph(
    "说明：本报告基于当前项目源码、Gradle 配置、AndroidManifest、页面资源与核心业务类的扫描结果整理生成，结论对应代码扫描时间为 2026 年 6 月 17 日。"
)

add_heading("一、项目概述", 1)
add_paragraph(
    "DayMark 是一款面向日常习惯培养场景的原生 Android 应用，围绕“创建习惯、设定提醒、每日打卡、查看统计、沉淀成长记录”这一主线展开。项目以本地数据存储为核心，不依赖服务器即可完成登录、习惯管理、图片记录、提醒通知、时间线回顾、周月统计、个人主页与勋章展示等功能，适合课程设计、作品演示以及离线使用场景。"
)
add_paragraph(
    "从当前代码实现来看，本项目不是简单的增删改查练习，而是已经形成了较完整的产品雏形：首页负责高频操作，编辑页承担习惯配置，日历页与时间线页负责历史回看，统计页承担数据可视化，个人主页负责数据汇总、头像资料与成就展示，整体模块边界比较清晰。"
)

add_heading("二、开发环境与技术栈", 1)
tech_table = doc.add_table(rows=1, cols=2)
tech_table.style = "Table Grid"
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
fill_cell(tech_table.rows[0].cells[0], "项目项", bold=True)
fill_cell(tech_table.rows[0].cells[1], "扫描结果", bold=True)
for left, right in [
    ("IDE", "Android Studio Panda 2 | 2025.3.2"),
    ("应用包名", "com.example.daymark"),
    ("构建环境", "Android Application，compileSdk 34，targetSdk 34，minSdk 23"),
    ("开发语言", "Java 11、XML 布局"),
    ("核心依赖", "Material 1.11.0、RecyclerView 1.3.2、Security Crypto 1.0.0"),
    ("本地存储", "SQLite、SharedPreferences、应用私有文件目录"),
    ("资源规模", "约 31 个 Java 文件、18 个布局 XML 文件"),
    ("关键系统能力", "相机、相册读取、通知权限、开机广播、FileProvider"),
]:
    row = tech_table.add_row().cells
    fill_cell(row[0], left, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(row[1], right, align=WD_ALIGN_PARAGRAPH.LEFT)

add_heading("三、需求完成情况", 1)
req_table = doc.add_table(rows=1, cols=3)
req_table.style = "Table Grid"
req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, text in enumerate(["题目要求", "代码实现情况", "完成度"]):
    fill_cell(req_table.rows[0].cells[i], text, bold=True)
for a, b, c in [
    ("登录功能，带记住密码", "由 LoginActivity 实现；结合 EncryptedSharedPreferences 保存记住密码状态、用户名与登录会话。", "已完成"),
    ("事件增、删、改", "MainActivity、EditHabitActivity 与 DayMarkDbHelper 共同完成新增、编辑、删除、排序等操作。", "已完成"),
    ("显示标题、内容、时间等", "首页列表与详情相关页面展示标题、内容、时间、分类、累计打卡次数等字段。", "已完成"),
    ("调用摄像头拍照、从相册选择图片", "编辑习惯页面支持相机拍照与相册选图，并通过 FileProvider / Uri 进行管理。", "已完成"),
    ("页面美观、交互较好", "项目采用 Material 风格、卡片式列表、折叠头部、时间线动画、热力图与勋章墙等设计。", "已完成"),
    ("功能可扩充", "当前已扩展模板库、提醒中心、日历、周月统计、个人主页、勋章、备份恢复等模块。", "已完成"),
]:
    row = req_table.add_row().cells
    fill_cell(row[0], a, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(row[1], b, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(row[2], c)

add_heading("四、系统功能设计", 1)
add_heading("4.1 登录与用户会话", 2)
add_paragraph(
    "登录模块由 LoginActivity 负责，提供用户名、密码输入和“记住密码”选项。项目已引入 androidx.security:security-crypto 依赖，并在登录与个人资料场景中使用 EncryptedSharedPreferences，能够避免将敏感信息直接以明文形式保存在普通偏好存储中。"
)
add_paragraph("数据库层同时提供注册和登录校验能力，说明项目结构已预留多用户扩展空间，而不是只写死单一演示账号。")

add_heading("4.2 首页与习惯列表管理", 2)
add_paragraph(
    "MainActivity 是项目的操作中枢。首页除显示习惯列表外，还提供“全部、今日待完成、今日已完成”三种过滤视图，支持列表排序与拖拽重排，并在上滑时收起部分按钮区、下滑回到顶部时重新展开，增强了移动端首页的信息密度和浏览效率。"
)
add_paragraph("首页还整合了新增习惯、个人主页、模板库、提醒中心、日历、统计、时间线以及导出入口，使主流程与扩展功能能够在一个页面完成跳转。")

add_heading("4.3 习惯新增、编辑与多维配置", 2)
add_paragraph(
    "EditHabitActivity 用于新增或编辑习惯，是题目要求中“事件增删改”的核心页面。页面支持录入标题、内容、时间、分类、目标天数、提醒时间，并提供多种频率配置方式，包括每日打卡、按周指定星期几、每周完成 N 次等。这样的设计让 APP 不仅能记录简单待办，也能覆盖阅读、跑步、背词等不同节律的习惯。"
)
add_paragraph("在媒体输入方面，编辑页支持调用系统相机拍照、从相册选图、即时预览图片，并把图片 Uri 与习惯数据一起保存，满足课程设计对相机与相册调用的明确要求。")

add_heading("4.4 打卡记录、日历与时间线", 2)
add_paragraph(
    "项目将“是否打卡成功”与“具体打卡记录”分开建模。首页的勾选与次数变化适合高频操作，而 CalendarActivity、DayDetailActivity、HistoryActivity 则面向回顾场景。日历页按日期展示打卡情况，支持回填历史某一天的完成记录；时间线页按记录倒序展示打卡轨迹，并提供按习惯筛选能力，适合答辩时演示“记录积累”的产品价值。"
)

add_heading("4.5 统计分析与提醒中心", 2)
add_paragraph("StatsActivity 负责周/月统计看板与分类占比可视化，帮助用户从“做了什么”进一步看到“坚持得怎么样”。ReminderCenterActivity 则集中显示已设置提醒的习惯，支持查看、编辑和关闭提醒，属于对基础打卡场景的实用增强。")

add_heading("4.6 个人主页、热力图与勋章墙", 2)
add_paragraph("ProfileActivity 提供头像、昵称、统计卡片、分类汇总、热力图、勋章墙、备份恢复、退出登录、删除账号等入口，已经不只是“个人资料页”，而是综合展示用户成长轨迹的中心页。AchievementsActivity 对全部勋章进行完整陈列；个人主页中的勋章墙使用不同样式展示勋章，增强了视觉层次与成就感。")

add_heading("4.7 模板库、导出与图片预览", 2)
add_paragraph("TemplateLibraryActivity 提供内置习惯模板，降低用户初次使用时的录入成本。MainActivity 中还包含导出文本报告的入口，ProfileActivity 中补充了 JSON 备份恢复能力。ImagePreviewActivity 则负责更好的图片查看体验，说明项目在“信息输入”和“信息查看”两个方向都做了相应收口。")

add_heading("五、核心模块实现分析", 1)
module_table = doc.add_table(rows=1, cols=3)
module_table.style = "Table Grid"
for i, text in enumerate(["核心类/页面", "职责", "说明"]):
    fill_cell(module_table.rows[0].cells[i], text, bold=True)
for a, b, c in [
    ("LoginActivity", "登录、记住密码、会话恢复", "与 EncryptedSharedPreferences 配合使用，承担应用入口。"),
    ("MainActivity", "首页列表与模块导航", "承担过滤、排序、重排、跳转、导出等高频操作。"),
    ("EditHabitActivity", "新增/编辑习惯", "覆盖标题、内容、时间、分类、频率、提醒、图片输入等字段。"),
    ("DayMarkDbHelper", "SQLite 数据管理", "负责建表、升级、事务、备份恢复、统计聚合，是项目核心数据层。"),
    ("HistoryActivity", "时间线展示", "按时间倒序展示打卡记录，并支持按习惯筛选。"),
    ("StatsActivity", "统计分析", "提供周/月统计与饼图分析。"),
    ("ProfileActivity", "个人中心", "负责头像昵称、热力图、勋章墙、备份恢复、退出登录等。"),
    ("ReminderReceiver / BootReceiver", "提醒广播", "支持定时提醒与重启后的提醒恢复。"),
    ("ImageUtils / ImageLoader", "图片处理", "负责图片 Uri、缓存、加载与文件处理。"),
]:
    row = module_table.add_row().cells
    fill_cell(row[0], a, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(row[1], b, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(row[2], c, align=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph("从结构上看，项目以 Activity + SQLiteOpenHelper + 自定义 View/Adapter 为主要组织方式，属于原生 Android 开发中较典型、也较适合课程设计验收的实现路径。各页面职责基本明确，便于演示时按业务流程逐步展开。")

add_heading("六、数据库设计", 1)
add_paragraph("项目数据库名称为 daymark.db，当前版本号为 7，在 onConfigure() 中启用了 WAL 模式，以提升并发读写体验。按照当前代码逻辑，数据库主要包含 users、habits、check_records 三张核心表。其关系可概括为：一个用户可以拥有多个习惯，一个习惯可以对应多条打卡记录。")

for title, rows in [
    (
        "6.1 users 表",
        [
            ("id", "INTEGER", "用户主键，自增"),
            ("username", "TEXT", "登录用户名，唯一"),
            ("password", "TEXT", "密码摘要或历史兼容字段"),
            ("display_name", "TEXT", "用户显示昵称"),
            ("salt", "TEXT", "密码加盐字段"),
            ("avatar_uri", "TEXT", "头像图片 Uri（后续版本已使用）"),
        ],
    ),
    (
        "6.2 habits 表",
        [
            ("id", "INTEGER", "习惯主键，自增"),
            ("user_id", "INTEGER", "所属用户"),
            ("title", "TEXT", "习惯标题"),
            ("content", "TEXT", "习惯内容说明"),
            ("time_text", "TEXT", "展示用时间文本"),
            ("image_uri", "TEXT", "习惯图片 Uri"),
            ("category", "TEXT", "分类"),
            ("reminder_time", "TEXT", "提醒时间"),
            ("check_count", "INTEGER", "累计打卡次数"),
            ("last_check_at", "INTEGER", "最近打卡时间戳"),
            ("created_at", "INTEGER", "创建时间"),
            ("frequency_type / frequency_days / frequency_count", "INTEGER / TEXT / INTEGER", "频率配置"),
            ("target_days", "INTEGER", "目标坚持天数"),
            ("sort_order", "INTEGER", "手动排序值"),
        ],
    ),
    (
        "6.3 check_records 表",
        [
            ("id", "INTEGER", "记录主键，自增"),
            ("habit_id", "INTEGER", "关联习惯 ID"),
            ("note", "TEXT", "补充说明或备注"),
            ("checked_at", "INTEGER", "打卡时间戳"),
        ],
    ),
]:
    add_heading(title, 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    fill_cell(table.rows[0].cells[0], "字段名", bold=True)
    fill_cell(table.rows[0].cells[1], "类型", bold=True)
    fill_cell(table.rows[0].cells[2], "说明", bold=True)
    for a, b, c in rows:
        row = table.add_row().cells
        fill_cell(row[0], a, align=WD_ALIGN_PARAGRAPH.LEFT)
        fill_cell(row[1], b, align=WD_ALIGN_PARAGRAPH.LEFT)
        fill_cell(row[2], c, align=WD_ALIGN_PARAGRAPH.LEFT)

add_heading("七、安全性与数据管理", 1)
add_paragraph("1. 登录信息保护：项目使用 EncryptedSharedPreferences 保存记住密码与会话信息，相比普通 SharedPreferences 更安全。")
add_paragraph("2. 密码处理：PasswordUtils 负责密码相关逻辑，数据库中存在 salt 字段，说明项目已朝加盐摘要方向演进。")
add_paragraph("3. 本地数据持久化：核心业务数据统一落在 SQLite 中，适合离线演示和无网络环境。")
add_paragraph("4. 备份与恢复：DayMarkDbHelper 提供 JSON 备份/恢复逻辑，ProfileActivity 中已接入导入恢复流程，有利于答辩时展示“数据可迁移”。")
add_paragraph("5. 图片访问与共享：应用使用 FileProvider 管理拍照结果文件，避免直接暴露文件路径。")

add_heading("八、界面与交互设计评价", 1)
add_paragraph("项目整体采用 Material 风格组件与卡片式布局，界面层级相对清晰。首页强调高频打卡，编辑页强调信息录入，统计页突出可视化，个人页聚焦成长反馈，这种按使用场景分层的方式比较符合真实产品的组织习惯。")
add_paragraph("从交互细节来看，项目已经实现了折叠按钮区、时间线入场动画、热力图、勋章墙、头像编辑、图片预览、列表拖拽排序等效果，较好地满足了题目中“页面美观、有良好交互性”的要求。对课程设计而言，这些设计既能提升演示观感，也能体现一定的产品思维。")

add_heading("九、构建、运行与测试情况", 1)
add_paragraph("1. 构建验证：在 2026 年 6 月 17 日的项目扫描中，执行 gradlew assembleDebug 可以成功通过，说明当前项目能够编译生成调试 APK。")
add_paragraph("2. APK 输出位置：app/build/outputs/apk/debug/app-debug.apk。")
add_paragraph("3. 静态检查：执行 gradlew lintDebug 时，Lint 报告指出 4 个错误、194 个警告；其中首个阻塞错误为主题资源循环（values-night-v27/themes.xml 中的 ResourceCycle）。这意味着项目功能层面可运行，但在资源组织和静态质量方面仍需继续收敛。")
add_paragraph("4. 手工测试覆盖面：从代码结构判断，项目已覆盖登录、习惯增删改查、打卡、图片输入、日历回填、提醒设置、时间线筛选、统计看板、头像与个人主页、备份恢复等主要用户路径。")

add_heading("十、项目亮点与创新点", 1)
add_paragraph("1. 在基础课程要求之外，项目扩展出了模板库、提醒中心、周/月统计、时间线、热力图、勋章墙、备份恢复等功能，完成度明显高于最低要求。")
add_paragraph("2. 习惯频率设计较完整，支持每日、每周指定日、每周次数等配置，较贴近真实习惯养成应用。")
add_paragraph("3. 首页、统计页、个人页分别服务于“执行”“复盘”“激励”三个不同阶段，产品思路较完整。")
add_paragraph("4. 项目以本地能力为主，演示稳定，不依赖网络环境，适合考试现场展示。")

add_heading("十一、当前不足与改进方向", 1)
issue_table = doc.add_table(rows=1, cols=4)
issue_table.style = "Table Grid"
for i, text in enumerate(["类别", "问题描述", "可能影响", "建议处理方向"]):
    fill_cell(issue_table.rows[0].cells[i], text, bold=True)
for a, b, c, d in [
    ("数据库一致性", "DayMarkDbHelper.onCreate() 当前创建 users 表时未直接包含 avatar_uri 列，但升级与资料模块已经依赖该列。", "新安装数据库在头像相关场景下存在 schema 不一致风险。", "统一 onCreate 与 onUpgrade 的最终 schema，确保 fresh install 与 upgrade install 行为一致。"),
    ("主题资源组织", "Lint 报告显示 values-night-v27/themes.xml 存在 Theme.DayMark / Theme.DayMark.Base 的资源循环。", "虽然 assembleDebug 可通过，但静态检查失败，影响项目质量与后续维护。", "重新整理 v27 与 night-v27 的主题继承链，只保留必要覆盖项。"),
    ("静态告警数量", "当前 lint 警告数量较多，说明资源、兼容性、代码规范层面仍有可清理空间。", "后续迭代时问题定位成本会提高。", "分批处理告警，优先消除阻塞错误和高频兼容性警告。"),
    ("产品延展性", "项目以本地单机为主，尚未引入云同步、桌面组件、跨设备协同等能力。", "长期使用与多端迁移体验有限。", "可继续扩展云备份、桌面小组件、搜索标签、成就分享等功能。"),
]:
    row = issue_table.add_row().cells
    fill_cell(row[0], a, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(row[1], b, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(row[2], c, align=WD_ALIGN_PARAGRAPH.LEFT)
    fill_cell(row[3], d, align=WD_ALIGN_PARAGRAPH.LEFT)

add_heading("十二、答辩演示建议", 1)
add_paragraph("建议按“登录 -> 新建习惯 -> 添加图片/提醒 -> 首页打卡 -> 时间线查看 -> 日历回看 -> 统计展示 -> 个人主页热力图与勋章墙 -> 备份/导出”这一顺序进行现场演示。这样既能覆盖课程要求，也能把扩展功能串成一条完整的使用闭环。")
add_paragraph("如果答辩时间较紧，可优先展示以下亮点：记住密码、频率配置、拍照/相册、提醒中心、周月统计、勋章墙、热力图和时间线。它们最能体现项目区别于普通待办清单的特色。")

add_heading("十三、总结", 1)
add_paragraph("综合本次扫描结果，DayMark 已经完成了“每日打卡习惯追踪 APP”课程设计的核心目标，并在登录安全、习惯配置、图片记录、提醒、统计、成长反馈等方面做出了较完整的产品化扩展。项目具备可编译、可演示、可答辩的基础，整体完成度较高。")
add_paragraph("后续若继续优化，建议优先修正数据库 schema 一致性与主题资源循环问题，再逐步清理 Lint 告警、增强搜索标签与数据同步能力。完成这些收尾后，项目的稳定性与正式感还会再上一个台阶。")

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
